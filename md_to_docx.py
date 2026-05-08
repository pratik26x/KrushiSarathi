"""
Convert PROJECT_REPORT.md to PROJECT_REPORT.docx (requires: pip install python-docx).
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "PROJECT_REPORT.md"
DOCX_PATH = ROOT / "PROJECT_REPORT.docx"


def add_runs_with_formatting(paragraph, text: str) -> None:
    """Split on **bold** and `code` segments."""
    if not text.strip():
        return
    # Split by **bold** or `code`
    pattern = r"(\*\*[^*]+\*\*|`[^`]+`)"
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if not line.startswith("|"):
        return []
    cells = [c.strip() for c in line.split("|")]
    if cells and cells[0] == "":
        cells = cells[1:]
    if cells and cells[-1] == "":
        cells = cells[:-1]
    return cells


def is_separator_row(cells: list[str]) -> bool:
    if not cells:
        return True
    return all(re.match(r"^[-:]+$", c.replace(" ", "")) for c in cells)


def main():
    text = MD_PATH.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = parse_table_row(lines[i])
                if cells and not is_separator_row(cells):
                    rows.append(cells)
                i += 1
            if rows:
                ncols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=ncols)
                table.style = "Table Grid"
                for ri, row_cells in enumerate(rows):
                    for ci in range(ncols):
                        cell_text = row_cells[ci] if ci < len(row_cells) else ""
                        cell = table.rows[ri].cells[ci]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        add_runs_with_formatting(p, cell_text)
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            title = stripped.lstrip("#").strip()
            if level == 1:
                doc.add_heading(title, level=1)
            elif level == 2:
                doc.add_heading(title, level=2)
            elif level == 3:
                doc.add_heading(title, level=3)
            else:
                doc.add_heading(title, level=4)
            i += 1
            continue

        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_formatting(p, re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_runs_with_formatting(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        if stripped == "":
            i += 1
            continue

        if stripped == "*End of detailed report.*":
            p = doc.add_paragraph()
            add_runs_with_formatting(p, stripped.strip("*"))
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs_with_formatting(p, line)
        i += 1

    doc.save(DOCX_PATH)
    print(f"Wrote {DOCX_PATH}")


if __name__ == "__main__":
    main()
